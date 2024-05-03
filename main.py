from docx import Document
from docx.shared import Pt


indexes = []
index = 0
answer = None
fields = [
    'Номер паспорта',
    'Фамилия',
    'Имя',
    'Имя отца',
    'Дата рождения',
    'Место рождения',
    'Тест',
    'Пол',
    'Дата выдачи',
    'Действителен до',
    'Длинный номер',
]
lines = []

# Открываем документ
doc = Document('template.docx')


def search_redact_paragraph(index):
    paragraph = doc.paragraphs[index]
    print([run.text for run in paragraph.runs])
    for run in paragraph.runs:
        run_text = run.text.strip()
        if run_text.startswith('<') and run_text.endswith('>'):
            indexes.append(index)

def apply_run_formatting(new_run, run):
    new_run.bold = run.bold
    new_run.italic = run.italic
    new_run.underline = run.underline
    new_run.font.size = run.font.size
    color = run.font.color.rgb
    if color is not None:
        new_run.font.color.rgb = color
    new_run.font.strike = run.font.strike
    return new_run

def redact_paragraph(index, line):
    # Получаем параграф
    flag = False # Показывает, изменялся ли уже текст в этом параграфе
    paragraph = doc.paragraphs[index]
    p_style = paragraph.style
    line_spacing = paragraph.paragraph_format.line_spacing
    p_alignment = paragraph.alignment

    # Создаем новый параграф
    new_paragraph = doc.add_paragraph(style=p_style)
    new_paragraph.paragraph_format.line_spacing = line_spacing
    new_paragraph.alignment = p_alignment

    # Проходимся по всем Run в текущем параграфе
    for run in paragraph.runs:
        run_text = run.text.strip()
        if (
            not run_text.startswith('<') and not run_text.endswith('>')
        ) or flag:
            # Если Run не содержит знаков < и > или flag == True,
            # то копируем текст Run без изменений, сохраняя оформление
            # Копируем Run в новый параграф
            new_run = new_paragraph.add_run(run.text)
            # Копируем форматирование
            new_run = apply_run_formatting(new_run, run)
            # print(new_run.font.size)
        else:
            # ситуация, когда в Run есть знаки < и >
            flag = True
            print(run.text)
            # new_run = new_paragraph.add_run()
            new_run = new_paragraph.add_run(line)
            # Копируем форматирование
            new_run = apply_run_formatting(new_run, run)
            if run.text[1:-1]:
                try:
                    f_s, new_run.bold = [
                        int(number) for number in run.text[1:-1].replace(
                            ' ', '').split(',')]
                    new_run.font.size = Pt(f_s)
                except ValueError:
                    array = [
                        int(number) for number in run.text[1:-1].replace(
                            ' ', ''
                            ).split(',')]
                    if len(array) != 2:
                        print(f'Ран {run.text} должен содержать '
                              f'два целых числа, разделённых запятой')

    # Вставляем новый параграф перед текущим
    doc.element.body.insert(index, new_paragraph._element)

    # Удаляем исходный параграф
    doc.element.body.remove(paragraph._element)



for field in fields:
    answer = input(f'{field}: ')
    if field in ('Дата рождения', 'Дата выдачи', 'Действителен до'):
        day, month, year = answer[:2], answer[2:4], answer[4:]
        answer = f'{day}/{month}/{year}'
    if field not in ('Пол', 'Место рождения'):
        answer = answer.upper()
    lines.append(answer)

print('----------------', lines)

for index in range(len(doc.paragraphs)):
    search_redact_paragraph(index)

print('------------', indexes)
for index, line in zip(indexes, lines):
    redact_paragraph(index, line)

print(list(zip(indexes, lines)))
# Сохраняем изменения
doc.save('test4.docx')


doc = Document('test4.docx')
for index in range(len(doc.paragraphs)):
    if search_redact_paragraph(index):
        indexes.append(index)
