import os

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from django.http import FileResponse


def add_run_to_paragraph(run, new_paragraph, part):
    # new_run = new_paragraph.add_run()
    new_run = new_paragraph.add_run(part)
    # Копируем форматирование
    new_run = apply_run_formatting(new_run, run)
    if run.text[1:-1]:
        try:
            f_s, new_run.bold = [
                int(number) for number in run.text[1:-1].replace(
                    ' ', '').split(',')]
            new_run.font.size = Pt(f_s)
        except ValueError:
            array = [
                int(number) for number in run.text[1:-1].replace(
                    ' ', ''
                ).split(',')]
            if len(array) != 2:
                print(f'Ран {run.text} должен содержать '
                        f'два целых числа, разделённых запятой')
    return new_run


def add_image_to_paragraph(paragraph, image_path):
    """
    Добавляет изображение в указанный параграф.

    :param paragraph: Параграф, в который нужно добавить изображение.
    :param image_path: Путь к изображению.
    """
    run = paragraph.add_run()

    # Добавляем изображение в параграф
    run.add_picture(image_path, width=Inches(6))  # Ширина изображения 6 дюймов

    return run



def search_redact_paragraph(doc, indexes, index):
    paragraph = doc.paragraphs[index]
    print(paragraph.text)
    for run in paragraph.runs:
        run_text = run.text.strip()
        if run_text.startswith('<') and run_text.endswith('>'):
            indexes.append(index)


def apply_run_formatting(new_run, run):
    new_run.bold = run.bold
    new_run.italic = run.italic
    new_run.underline = run.underline
    new_run.font.size = run.font.size
    color = run.font.color.rgb
    if color is not None:
        new_run.font.color.rgb = color
    new_run.font.strike = run.font.strike
    return new_run


def redact_paragraph(doc, index, line, cleaned_data):
    # Получаем параграф
    flag = False  # Показывает, изменялся ли уже текст в этом параграфе
    paragraph = doc.paragraphs[index]
    p_style = paragraph.style
    line_spacing = paragraph.paragraph_format.line_spacing
    p_alignment = paragraph.alignment

    # Создаем новый параграф
    new_paragraph = doc.add_paragraph(style=p_style)
    new_paragraph.paragraph_format.line_spacing = line_spacing
    new_paragraph.alignment = p_alignment

    # Проходимся по всем Run в текущем параграфе
    for run in paragraph.runs:
        run_text = run.text.strip()
        if (
            not run_text.startswith('<') and not run_text.endswith('>')
        ) or flag:
            # Если Run не содержит знаков < и > или flag == True,
            # то копируем текст Run без изменений, сохраняя оформление
            # Копируем Run в новый параграф
            new_run = new_paragraph.add_run(run.text)
            # Копируем форматирование
            new_run = apply_run_formatting(new_run, run)
            # print(new_run.font.size)
        else:
            # ситуация, когда в Run есть знаки < и >
            flag = True
            print(run.text)
            if run.text[1:-1] == 'image':
                if cleaned_data['image'] is not None:
                    add_image_to_paragraph(new_paragraph, cleaned_data['image'])  # cleaned_data['image'] - объект изображения, загруженный через форму
            else:
                lines = line.split('\n')
                if len(lines) > 1:
                    for i in range(len(lines) - 1): # Обрабатываем вариант, когда значение многострочное
                        new_run = add_run_to_paragraph(run, new_paragraph, lines[i])
                        new_run.add_break()
                new_run = add_run_to_paragraph(run, new_paragraph, lines[-1])

    # Вставляем новый параграф перед текущим
    doc.element.body.insert(index, new_paragraph._element)

    # Удаляем исходный параграф
    doc.element.body.remove(paragraph._element)


def main(cleaned_data, image_path=None):
    # Определите путь к файлу template.docx относительно текущего файла main.py
    current_dir = os.path.dirname(__file__)
    template_path = os.path.join(current_dir, 'template.docx')
    # Теперь используйте template_path для загрузки документа
    doc = Document(template_path)
    indexes = []
    index = 0
    answer = None
    lines = []
    print('Мы в функции main файла main.py')
    file_name = cleaned_data.pop('file_name')
    print('file_name =', file_name)
    if 'image' in cleaned_data:
        clean_d = [('image', cleaned_data['image'])] + list(cleaned_data.items())[:-1]
        print('\n------\n\n', clean_d, '\n--------\n')
    for field_name, data in clean_d:
        answer = data
        if field_name in ('birthday', 'date_of_issue', 'date_of_expiry'):
            day, month, year = answer[:2], answer[2:4], answer[4:]
            answer = f'{day}/{month}/{year}'
        if field_name not in ('place_of_birthday', 'sex', 'image', 'issuing_authority', 'signature'):
            answer = answer.upper()
        if field_name == 'place_of_birthday':
            answer = answer.capitalize()
        if field_name == 'sex':
            answer = {'M': 'мужской', 'F': 'женский'}[answer]
        if field_name == 'issuing_authority':
            answer = {"authority1": "Бригадный генерал\nАли Золгадри",
                      "authority2": "Бригадный генерал\nСадэг Резадуст"}[answer]
        if field_name == 'signature':
            answer = {"empty": "-",
                      "signature": "/подпись/"}[answer]

        lines.append(answer)

    print('----------------', lines)

    for index in range(len(doc.paragraphs)):
        search_redact_paragraph(doc, indexes, index)

    print('------------\n', indexes)
    for index, line in zip(indexes, lines):
        redact_paragraph(doc, index, line, cleaned_data)

    print(list(zip(indexes, lines)))

#     # Определите путь для сохранения файла
#     output_path = os.path.join(current_dir, f'{file_name}.docx')
# # Сохраняем изменения
#     doc.save(output_path)

#     doc.save(f'{file_name}.docx')
    output_stream = BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream, file_name


if __name__ == '__main__':
    main()


def custom_filter(some_list):
    return sum(filter(lambda x: not x % 7, filter(
        lambda x: isinstance(x, int), some_list))) <= 83