import os

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def runs_are_similar(run1, run2):
    """Проверяет, имеют ли два рана одинаковое оформление."""
    return (
        run1.bold == run2.bold and
        run1.italic == run2.italic and
        run1.underline == run2.underline and
        run1.font.size == run2.font.size and
        run1.font.color.rgb == run2.font.color.rgb and
        run1.font.strike == run2.font.strike
    )


def merge_similar_runs(doc):
    for para in doc.paragraphs:
        if '<w:br w:type="page"/>' in para._element.xml:
            continue
        if len(para.runs) < 2:
            continue
        new_runs = []
        current_run = para.runs[0]
        current_text = current_run.text
        for run in para.runs[1:]:
            if '<w:br w:type="page"/>' in run._element.xml:
                print("Run содержит разрыв страницы")

            if runs_are_similar(current_run, run):
                current_text += run.text
            else:
                new_run = para.add_run(current_text)
                new_run.bold = current_run.bold
                new_run.italic = current_run.italic
                new_run.underline = current_run.underline
                new_run.font.size = current_run.font.size
                color = current_run.font.color.rgb
                if color is not None:
                    new_run.font.color.rgb = color
                new_run.font.strike = current_run.font.strike

                new_runs.append(new_run)
                current_run = run
                current_text = run.text

        new_run = para.add_run(current_text)
        new_run.bold = current_run.bold
        new_run.italic = current_run.italic
        new_run.underline = current_run.underline
        new_run.font.size = current_run.font.size
        color = current_run.font.color.rgb
        if color is not None:
            new_run.font.color.rgb = color
        new_run.font.strike = current_run.font.strike

        new_runs.append(new_run)

        # Remove the old runs and add the merged ones
        for run in para.runs:
            run._element.getparent().remove(run._element)

        for new_run in new_runs:
            para._element.append(new_run._element)


# Определяем маркеры полей
field_start_marker = '<<'
field_end_marker = '>>'


def find_and_process_markers(doc):
    fields = []
    in_field = False
    field_content = []

    for para in doc.paragraphs:
        for run in para.runs:
            text = run.text
            start_pos = text.find(field_start_marker)
            end_pos = text.find(field_end_marker)

            if start_pos != -1:
                in_field = True
                if end_pos != -1:
                    # Если оба маркера в одном ране
                    fields.append(text[start_pos + len(field_start_marker):end_pos])
                    run.text = text[:start_pos] + text[start_pos + len(field_start_marker):end_pos] + text[end_pos + len(field_end_marker):]
                    in_field = False
                else:
                    # Если только начальный маркер в ране
                    field_content.append(text[start_pos + len(field_start_marker):])
                    run.text = text[:start_pos]

            elif end_pos != -1 and in_field:
                # Если конечный маркер в ране
                field_content.append(text[:end_pos])
                fields.append(''.join(field_content))
                run.text = ''.join(field_content) + text[end_pos + len(field_end_marker):]
                field_content = []
                in_field = False

            elif in_field:
                # Если внутри поля
                field_content.append(text)
                run.text = ''.join(field_content)

    return fields


def process_docx(current_dir, file_name):
    file_path = os.path.join(current_dir, file_name)
    doc = Document(file_path)
    file_path = os.path.join(current_dir, 'temp.docx')
    merge_similar_runs(doc)
    doc.save(file_path)
    doc = Document(file_path)
    fields = find_and_process_markers(doc)
    file_path = os.path.join(current_dir, 'processed_.docx')
    doc.save(file_path)
    return fields


# Пример использования
current_dir = os.path.dirname(__file__)
  # Укажите путь к вашему шаблону
fields = process_docx(current_dir, 'raw_template.docx')
print(fields)
